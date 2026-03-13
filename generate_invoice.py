"""
Boring Science LLC — Invoice DOCX Generator
Page 1: White background, print-safe invoice
Page 2: Dark brand page with logo, taglines, services

Run: python3 generate_invoice.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── COLOURS ─────────────────────────────────────────────────────
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x0A, 0x0A, 0x0A)
LIME   = RGBColor(0xC8, 0xFF, 0x00)   # brand accent
DARK   = RGBColor(0x0A, 0x0A, 0x0A)   # page 2 background
MUTED  = RGBColor(0x88, 0x88, 0x88)   # grey labels
LIGHT  = RGBColor(0xF5, 0xF5, 0xF5)   # very light bg (page 1 bands)
BORDER = RGBColor(0xDD, 0xDD, 0xDD)   # light border (page 1)
DIM    = RGBColor(0xBB, 0xBB, 0xBB)   # dimmed white (page 2)
FAINT  = RGBColor(0x44, 0x44, 0x44)   # barely visible on dark (dividers)

MONO  = "Courier New"
SANS  = "Calibri"
SERIF = "Georgia"

# ── HELPERS ─────────────────────────────────────────────────────
def set_cell_bg(cell, rgb):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  str(rgb))
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr    = cell._tc.get_or_add_tcPr()
    tcBords = OxmlElement('w:tcBorders')
    for side, spec in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if spec:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   spec.get('val', 'single'))
            el.set(qn('w:sz'),    str(spec.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), spec.get('color', '000000'))
            tcBords.append(el)
    tcPr.append(tcBords)

def no_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    bdr = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        bdr.append(el)
    tblPr.append(bdr)

def run(para, text, bold=False, italic=False, font=SANS, size=10, color=BLACK):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = font; r.font.size = Pt(size); r.font.color.rgb = color
    return r

def lbl(para, text, color=MUTED):
    para.clear()
    r = para.add_run(text.upper())
    r.font.name = MONO; r.font.size = Pt(7); r.font.color.rgb = color
    para.paragraph_format.space_after = Pt(2)

def hrule(doc, color=BORDER):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1');      bot.set(qn('w:color'), str(color))
    pBdr.append(bot); pPr.append(pBdr)

def gap(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# DOCUMENT SETUP
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
doc = Document()
sec = doc.sections[0]
sec.page_width    = Cm(21.0);  sec.page_height   = Cm(29.7)
sec.top_margin    = Cm(1.5);   sec.bottom_margin = Cm(1.5)
sec.left_margin   = Cm(1.8);   sec.right_margin  = Cm(1.8)
CONTENT_W = Cm(17.4)   # 21 - 1.8 - 1.8

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PAGE 1 — INVOICE (WHITE, PRINT-SAFE)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# ── HEADER BAR (lime green) ──────────────────────────────────────
ht = doc.add_table(rows=1, cols=2)
no_borders(ht); ht.alignment = WD_TABLE_ALIGNMENT.CENTER
lc, rc = ht.cell(0,0), ht.cell(0,1)
lc.width = Cm(11.4); rc.width = Cm(6.0)
for c in [lc, rc]:
    set_cell_bg(c, LIME)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Brand — left
bp = lc.add_paragraph()
bp.paragraph_format.space_before = Pt(10); bp.paragraph_format.space_after = Pt(2)
run(bp, 'BORING SCIENCE', bold=True, font=MONO, size=18, color=BLACK)
sp = lc.add_paragraph()
sp.paragraph_format.space_after = Pt(10)
run(sp, 'boringscience.bio  ·  Sunnyvale, California', font=MONO, size=7.5, color=BLACK)

# Invoice number — right
ip = rc.add_paragraph()
ip.alignment = WD_ALIGN_PARAGRAPH.RIGHT
ip.paragraph_format.space_before = Pt(10); ip.paragraph_format.space_after = Pt(3)
run(ip, 'INVOICE', font=MONO, size=7.5, color=BLACK)
np_ = rc.add_paragraph()
np_.alignment = WD_ALIGN_PARAGRAPH.RIGHT; np_.paragraph_format.space_after = Pt(10)
run(np_, '#INV-2026-001', bold=True, font=MONO, size=16, color=BLACK)

gap(doc, 8)

# ── FROM / BILL TO ───────────────────────────────────────────────
pt = doc.add_table(rows=1, cols=2)
no_borders(pt); pt.alignment = WD_TABLE_ALIGNMENT.CENTER
fc, tc_ = pt.cell(0,0), pt.cell(0,1)
fc.width = Cm(8.7); tc_.width = Cm(8.7)

for cell, title, lines in [
    (fc,  'From',    ['Boring Science LLC', '1234 Innovation Drive, Suite 100',
                      'Sunnyvale, CA 94086  ·  United States',
                      'inquiries@boringscience.bio  ·  408.368.9547']),
    (tc_, 'Bill To', ['Client Name / Organization', 'Street Address',
                      'City, State ZIP  ·  Country', 'client@example.com']),
]:
    set_cell_bg(cell, WHITE)
    lp = cell.add_paragraph(); lbl(lp, title)
    for i, line in enumerate(lines):
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(1)
        run(p2, line, font=SANS, size=9.5 if i == 0 else 9,
            color=BLACK if i == 0 else MUTED, bold=(i == 0))

gap(doc, 8)

# ── DATES ROW ────────────────────────────────────────────────────
dt = doc.add_table(rows=1, cols=4)
no_borders(dt); dt.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (title, value) in enumerate([
    ('Issue Date',       '2026-03-12'),
    ('Due Date',         '2026-04-11'),
    ('Project / PO Ref', 'BS-PROJECT-001'),
    ('Terms',            'Net 30'),
]):
    c = dt.cell(0, i); c.width = Cm(4.35)
    set_cell_bg(c, LIGHT)
    set_cell_borders(c,
        top   ={'sz':4,'color':str(BORDER)},
        bottom={'sz':4,'color':str(BORDER)},
        left  ={'sz':4,'color':str(BORDER)},
        right ={'sz':4,'color':str(BORDER)})
    lp2 = c.add_paragraph(); lbl(lp2, title)
    lp2.paragraph_format.space_before = Pt(6)
    vp = c.add_paragraph(); vp.paragraph_format.space_after = Pt(6)
    run(vp, value, font=MONO, size=9, color=BLACK, bold=(title == 'Due Date'))

gap(doc, 8)

# ── LINE ITEMS ───────────────────────────────────────────────────
lbl_si = doc.add_paragraph(); lbl(lbl_si, 'Services Rendered')

it = doc.add_table(rows=1, cols=4)
no_borders(it); it.alignment = WD_TABLE_ALIGNMENT.CENTER
col_w = [9.8, 1.8, 2.9, 2.9]

# Header row
for i, (hdr, w) in enumerate(zip(['Description','Qty','Rate (USD)','Amount (USD)'], col_w)):
    c = it.rows[0].cells[i]; c.width = Cm(w)
    set_cell_bg(c, LIGHT)
    set_cell_borders(c,
        top   ={'sz':4,'color':str(BORDER)},
        bottom={'sz':6,'color':str(BLACK)})
    hp = c.add_paragraph()
    hp.paragraph_format.space_before = Pt(4); hp.paragraph_format.space_after = Pt(4)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
    run(hp, hdr.upper(), font=MONO, size=7, color=MUTED)

# Data rows
for ri, (desc, sub, qty, rate, amt) in enumerate([
    ('Computational Biology Consulting',
     'SE(3)-GNN docking pipeline setup and validation',
     '16 hrs', '$250.00', '$4,000.00'),
    ('NGS Pipeline Engineering',
     'Nextflow DSL2 WGS pipeline — containerization & CI',
     '24 hrs', '$250.00', '$6,000.00'),
    ('bsdock License — Annual',
     'Single-site, unlimited users. Includes support & updates.',
     '1', '$3,500.00', '$3,500.00'),
    ('Data Report & Documentation',
     'Technical report, methodology writeup, reproducibility pkg',
     '8 hrs', '$200.00', '$1,600.00'),
]):
    row = it.add_row()
    bg  = LIGHT if ri % 2 == 0 else WHITE
    for i, c in enumerate(row.cells):
        c.width = Cm(col_w[i])
        set_cell_bg(c, bg)
        set_cell_borders(c, bottom={'sz':2,'color':str(BORDER)})
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    dp = row.cells[0].add_paragraph()
    dp.paragraph_format.space_before = Pt(5); dp.paragraph_format.space_after = Pt(1)
    run(dp, desc, font=SANS, size=9.5, color=BLACK, bold=True)
    sbp = row.cells[0].add_paragraph(); sbp.paragraph_format.space_after = Pt(5)
    run(sbp, sub, font=SANS, size=8, color=MUTED)

    for ci, (val, bold_it) in enumerate([(qty, False),(rate, False),(amt, True)], 1):
        vp2 = row.cells[ci].add_paragraph()
        vp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        vp2.paragraph_format.space_before = Pt(5); vp2.paragraph_format.space_after = Pt(5)
        run(vp2, val, font=MONO, size=9, color=BLACK, bold=bold_it)

gap(doc, 4)

# ── TOTALS ───────────────────────────────────────────────────────
tot = doc.add_table(rows=4, cols=2)
no_borders(tot); tot.alignment = WD_TABLE_ALIGNMENT.RIGHT

for i, (title, value, bold_it) in enumerate([
    ('Subtotal',  '$15,100.00', False),
    ('Tax (0%)',  '$0.00',      False),
    ('Discount',  '—',          False),
    ('Total Due', '$15,100.00', True),
]):
    lc2, vc = tot.rows[i].cells[0], tot.rows[i].cells[1]
    lc2.width = Cm(4.0); vc.width = Cm(3.5)
    set_cell_bg(lc2, WHITE); set_cell_bg(vc, WHITE)

    if i == 3:
        set_cell_borders(lc2, top={'sz':4,'color':str(BORDER)})
        set_cell_borders(vc,  top={'sz':4,'color':str(BORDER)})

    lp3 = lc2.add_paragraph(); lp3.paragraph_format.space_after = Pt(2)
    run(lp3, title.upper(), font=MONO, size=7.5, color=BLACK if bold_it else MUTED, bold=bold_it)

    vp3 = vc.add_paragraph()
    vp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT; vp3.paragraph_format.space_after = Pt(2)
    run(vp3, value, font=MONO, size=13 if bold_it else 9, color=BLACK, bold=bold_it)

gap(doc, 6)
hrule(doc, BORDER)

# ── PAYMENT INSTRUCTIONS ─────────────────────────────────────────
pay = doc.add_table(rows=1, cols=1)
no_borders(pay); pay.alignment = WD_TABLE_ALIGNMENT.CENTER
pc = pay.cell(0, 0); pc.width = CONTENT_W
set_cell_bg(pc, LIGHT)
set_cell_borders(pc,
    top   ={'sz':4,'color':str(BORDER)},
    bottom={'sz':4,'color':str(BORDER)},
    right ={'sz':4,'color':str(BORDER)},
    left  ={'sz':18,'color':str(LIME)})   # thick lime left accent

lp4 = pc.add_paragraph(); lbl(lp4, 'Payment Instructions')
lp4.paragraph_format.space_before = Pt(8)

pgrid = pc.add_table(rows=3, cols=4)
no_borders(pgrid)
pay_items = [
    ('Bank',          'Bank Name'),
    ('Account Name',  'Boring Science LLC'),
    ('Account No.',   'XXXX-XXXX-XXXX'),
    ('Routing/SWIFT', 'XXXXXXXXX'),
    ('Wire Reference','#INV-2026-001'),
    ('Also Accepts',  'ACH · Check · Stripe'),
]
for ri in range(3):
    for ci in range(4):
        idx = ri * 2 + ci // 2
        c   = pgrid.rows[ri].cells[ci]
        set_cell_bg(c, LIGHT)
        if idx < len(pay_items):
            if ci % 2 == 0:
                lp5 = c.add_paragraph(); lp5.paragraph_format.space_after = Pt(1)
                run(lp5, pay_items[idx][0].upper(), font=MONO, size=7, color=MUTED)
            else:
                vp4 = c.add_paragraph(); vp4.paragraph_format.space_after = Pt(4)
                run(vp4, pay_items[idx][1], font=MONO, size=9, color=BLACK)
ep = pc.add_paragraph(); ep.paragraph_format.space_after = Pt(6)

gap(doc, 6)

# ── NOTES ────────────────────────────────────────────────────────
lbl_n = doc.add_paragraph(); lbl(lbl_n, 'Notes')
np2 = doc.add_paragraph(); np2.paragraph_format.space_after = Pt(8)
run(np2,
    'Payment is due within 30 days of the invoice date. '
    'Late payments are subject to a 1.5% monthly interest charge. '
    'Please include the invoice number in your payment reference. '
    'Questions: inquiries@boringscience.bio',
    font=SANS, size=8.5, color=MUTED)

hrule(doc, BORDER)

# ── PAGE 1 FOOTER ────────────────────────────────────────────────
ft = doc.add_table(rows=1, cols=2)
no_borders(ft); ft.alignment = WD_TABLE_ALIGNMENT.CENTER
fl, fr = ft.cell(0,0), ft.cell(0,1)
fl.width = Cm(9.5); fr.width = Cm(7.9)
set_cell_bg(fl, WHITE); set_cell_bg(fr, WHITE)
flp = fl.add_paragraph()
run(flp, '© 2026 BORING SCIENCE LLC  ·  boringscience.bio  ·  408.368.9547',
    font=MONO, size=7, color=MUTED)
frp = fr.add_paragraph(); frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run(frp, "Science is unpredictable. Your invoices shouldn't be.",
    font=SERIF, size=9, color=MUTED, italic=True)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PAGE 2 — DARK BRAND PAGE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
doc.add_page_break()

# Full-width dark container
brand = doc.add_table(rows=1, cols=1)
no_borders(brand); brand.alignment = WD_TABLE_ALIGNMENT.CENTER
bc = brand.cell(0, 0); bc.width = CONTENT_W
set_cell_bg(bc, DARK)

def dp(text='', align=WD_ALIGN_PARAGRAPH.CENTER, after=8, before=0,
       font=SANS, size=10, color=DIM, bold=False, italic=False):
    """Quick dark-page paragraph"""
    p = bc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if text:
        r = p.add_run(text)
        r.font.name = font; r.font.size = Pt(size)
        r.font.color.rgb = color; r.bold = bold; r.italic = italic
    return p

def dline():
    """Subtle divider on dark page"""
    p = bc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run('─' * 48); r.font.name = MONO
    r.font.size = Pt(8); r.font.color.rgb = FAINT

# Top spacer
for _ in range(4): dp(after=10)

# ── WORDMARK ─────────────────────────────────────────────────────
wmp = bc.add_paragraph()
wmp.alignment = WD_ALIGN_PARAGRAPH.CENTER
wmp.paragraph_format.space_before = Pt(0)
wmp.paragraph_format.space_after  = Pt(6)
r1 = wmp.add_run('BORING ')
r1.font.name = MONO; r1.font.size = Pt(40); r1.font.color.rgb = WHITE; r1.bold = True
r2 = wmp.add_run('SCIENCE')
r2.font.name = MONO; r2.font.size = Pt(40); r2.font.color.rgb = LIME;  r2.bold = True

# ── PRIMARY TAGLINE ───────────────────────────────────────────────
dp("Science is unpredictable. Your data shouldn't be.",
   font=SERIF, size=14, color=DIM, italic=True, after=32)

dline()

# ── SERVICES ─────────────────────────────────────────────────────
dp('WHAT WE BUILD', font=MONO, size=8, color=LIME, after=16)

for svc, desc in [
    ('Computational Biology',   'Protein structure · Molecular docking · ML/AI pipelines'),
    ('NGS & Genomics',          'WGS · RNA-seq · ChIP-seq · Variant calling · GWAS'),
    ('Data Infrastructure',     'Nextflow · Snakemake · Cloud HPC · Reproducibility'),
    ('Software & Licensing',    'bsdock · boringlab · Custom bioinformatics tooling'),
]:
    p = bc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    rs = p.add_run(svc.upper() + '   ')
    rs.font.name = MONO; rs.font.size = Pt(9); rs.font.color.rgb = LIME; rs.bold = True
    rd = p.add_run(desc)
    rd.font.name = SANS; rd.font.size = Pt(9); rd.font.color.rgb = DIM

# ── ABOUT ────────────────────────────────────────────────────────
dp(after=24)
dline()

dp('Boring Science builds the rigorous data infrastructure that turns data points into',
   font=SANS, size=10.5, color=DIM, after=4)
dp('discovery — clean, structured, reproducible, and audit-ready.',
   font=SANS, size=10.5, color=DIM, after=28)

dline()

# ── CONTACT ──────────────────────────────────────────────────────
dp('GET IN TOUCH', font=MONO, size=8, color=LIME, after=14)

for line, col in [
    ('boringscience.bio',          WHITE),
    ('inquiries@boringscience.bio', DIM),
    ('408.368.9547',               DIM),
    ('Sunnyvale, CA  ·  Ottawa, ON', DIM),
]:
    dp(line, font=MONO, size=10, color=col, after=6)

# Bottom spacer
for _ in range(3): dp(after=12)

# Legal
dp('© 2026 BORING SCIENCE LLC  ·  ALL RIGHTS RESERVED',
   font=MONO, size=7, color=FAINT, after=10)

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SAVE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
doc.save('invoice-template.docx')
print('✓  Saved: invoice-template.docx')
